import csv
import json
from io import BytesIO, StringIO
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from pypdf import PdfReader, PdfWriter

from .models import Evidence


COLUMNS = [
    ("index", "序号"), ("company", "企业名称"), ("indicator", "指标"),
    ("value_candidate", "数值候选（待复核）"), ("currency", "币种（原文）"),
    ("amount_unit", "金额单位（原文）"), ("value_year", "数值所属年份"),
    ("source_filename", "来源文件名"), ("page_number", "页码"),
    ("language", "语种"), ("year", "报告年份"),
    ("document_type", "原文件类型"), ("matched_terms", "命中词"),
    ("score", "匹配分"), ("excerpt", "证据片段"),
    ("evidence_filename", "证据文件名"), ("review_status", "复核状态"),
]


def extract_single_page(pdf_bytes: bytes, page_number: int) -> bytes:
    reader = PdfReader(BytesIO(pdf_bytes))
    if page_number < 1 or page_number > len(reader.pages):
        raise ValueError(f"页码超出范围：{page_number}")
    writer = PdfWriter()
    writer.add_page(reader.pages[page_number - 1])
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


def build_excel(results: list[Evidence]) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "提取结果"
    ws.append([label for _, label in COLUMNS])
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E78")
        cell.alignment = Alignment(horizontal="center", vertical="center")
    for result in results:
        record = result.to_record()
        ws.append([record[key] for key, _ in COLUMNS])
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    widths = [8, 22, 18, 22, 16, 18, 14, 28, 8, 10, 12, 16, 24, 10, 70, 42, 14]
    for idx, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    output = BytesIO()
    wb.save(output)
    return output.getvalue()


def build_word(results: list[Evidence]) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.8)
    section.left_margin = section.right_margin = Inches(0.9)

    def set_font(run, size=10, bold=False, color=None):
        run.font.name = "Arial Unicode MS"
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial Unicode MS")
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")

    normal = doc.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal.font.size = Pt(10)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    set_font(title.add_run("PDF 证据提取结果"), size=22, bold=True, color=(31, 78, 121))
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(14)
    set_font(note.add_run("说明：自动识别结果仅用于辅助定位。数值、币种、单位和证据原页均应人工复核。"), size=9, color=(89, 89, 89))

    for result in results:
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(6)
        set_font(
            heading.add_run(f"{result.index:03d}  {result.company} - 第 {result.page_number} 页"),
            size=14, bold=True, color=(31, 78, 121),
        )
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        items = [
            ("指标", result.indicator), ("数值候选", result.value_candidate),
            ("币种（原文）", result.currency), ("金额单位（原文）", result.amount_unit),
            ("数值所属年份", result.value_year),
            ("来源文件", result.source_filename), ("证据文件", result.evidence_filename),
            ("命中词", result.matched_terms), ("匹配分", str(result.score)),
            ("证据片段", result.excerpt),
        ]
        for label, value in items:
            cells = table.add_row().cells
            cells[0].width, cells[1].width = Inches(1.35), Inches(5.25)
            for cell in cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                tc_mar = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcMar")
                if tc_mar is None:
                    tc_mar = OxmlElement("w:tcMar")
                    cell._tc.get_or_add_tcPr().append(tc_mar)
                for edge in ("top", "start", "bottom", "end"):
                    node = OxmlElement(f"w:{edge}")
                    node.set(qn("w:w"), "100")
                    node.set(qn("w:type"), "dxa")
                    tc_mar.append(node)
            shade = OxmlElement("w:shd")
            shade.set(qn("w:fill"), "E8EEF5")
            cells[0]._tc.get_or_add_tcPr().append(shade)
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_font(cells[0].paragraphs[0].add_run(label), size=9, bold=True, color=(31, 78, 121))
            set_font(cells[1].paragraphs[0].add_run(str(value)), size=9)
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def build_csv(results: list[Evidence]) -> bytes:
    output = StringIO()
    writer = csv.writer(output)
    writer.writerow([label for _, label in COLUMNS])
    for result in results:
        record = result.to_record()
        writer.writerow([record[key] for key, _ in COLUMNS])
    return output.getvalue().encode("utf-8-sig")


def build_export_bundle(results: list[Evidence], source_pdfs: dict[str, bytes]) -> dict[str, bytes]:
    files = {
        "提取结果.xlsx": build_excel(results),
        "证据清单.docx": build_word(results),
        "提取结果.csv": build_csv(results),
    }
    for result in results:
        files[f"证据页/{result.evidence_filename}"] = extract_single_page(
            source_pdfs[result.source_filename], result.page_number
        )
    manifest = {
        "result_count": len(results),
        "notice": "自动提取结果仅供辅助定位，请人工复核证据原页。",
        "files": [result.to_record() for result in results],
    }
    files["manifest.json"] = json.dumps(manifest, ensure_ascii=False, indent=2).encode("utf-8")
    zip_buffer = BytesIO()
    with ZipFile(zip_buffer, "w", ZIP_DEFLATED) as archive:
        for name, data in files.items():
            archive.writestr(name, data)
    files["证据提取结果包.zip"] = zip_buffer.getvalue()
    return files
